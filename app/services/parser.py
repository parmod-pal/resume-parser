import re
import string
import fitz  # PyMuPDF
from docx import Document
from app.utils.logger import logger

def extract_text(file_path: str, filename: str) -> str:
    """Extract raw text from PDF, DOCX, or DOC, including tables."""
    try:
        text = ""
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            with fitz.open(file_path) as doc:
                text = "\n".join([page.get_text("text") for page in doc])
                
        elif filename_lower.endswith('.docx'):
            doc = Document(file_path)
            extracted_lines = []
            
            # 1. Extract standard paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    extracted_lines.append(p.text.strip())
                    
            # 2. Extract tables (Crucial for resumes!)
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text.strip() and cell.text.strip() not in row_data:
                            row_data.append(cell.text.strip())
                    if row_data:
                        extracted_lines.append(" | ".join(row_data))
                        
            text = "\n".join(extracted_lines)
            
        elif filename_lower.endswith('.doc'):
            # Fallback binary string extraction for legacy .doc files
            with open(file_path, "rb") as f:
                content = f.read()
            printable = set(bytes(string.printable, 'ascii'))
            text = ''.join([chr(c) if c in printable else '\n' for c in content])
            
        else:
            raise ValueError("Unsupported file type")
        
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
    except Exception as e:
        logger.error(f"Text extraction failed: {e}")
        raise

def extract_rules_based_data(text: str) -> dict:
    """Extract deterministic data using Regex."""
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    phone_pattern = r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    url_pattern = r'(https?://[^\s]+|www\.[^\s]+|linkedin\.com/in/[^\s]+|github\.com/[^\s]+)'
    
    emails = re.findall(email_pattern, text)
    phones = re.findall(phone_pattern, text)
    urls = re.findall(url_pattern, text)
    
    return {
        "email": emails[0] if emails else None,
        "phone": phones[0] if phones else None,
        "urls": list(set(urls))
    }